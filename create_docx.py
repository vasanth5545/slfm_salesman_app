from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# Name
name = doc.add_heading('VASANTH M', 0)
name.alignment = 1 # Center

# Contact info
contact = doc.add_paragraph('41/12A, MaligaiMedu, Mannargudi, Thiruvarur District, Tamil Nadu\nPhone: +91 8489758648 | Email: vasanthvarman0@gmail.com | GitHub: github.com/vasanthvarman0')
contact.alignment = 1

doc.add_heading('Professional Profile', level=1)
doc.add_paragraph('A highly motivated and results-oriented App Developer and Visual Designer with hands-on experience in building robust Flutter-based mobile applications, developing custom PHP backend APIs, and crafting impactful visuals. Proven ability to architect complex, offline-first mobile solutions with enterprise-grade security and dynamic server-driven functionalities. Eager to contribute technical expertise and creative problem-solving skills to innovative projects.')

doc.add_heading('Technical Skills', level=1)
doc.add_paragraph('Mobile App Development: Flutter, Dart, Offline-First Architecture, Server-Driven UI.', style='List Bullet')
doc.add_paragraph('Backend & APIs: PHP (Custom API Development), REST APIs (Dio, http), MySQL.', style='List Bullet')
doc.add_paragraph('State Management & Local Storage: Provider, SQLite, Hive, Flutter Secure Storage.', style='List Bullet')
doc.add_paragraph('Cloud Integration: Firebase (Firestore, Realtime DB, Remote Config, Cloud Messaging, Crashlytics).', style='List Bullet')
doc.add_paragraph('Security & Hardware APIs: FreeRASP (Root/Jailbreak detection), SSL Pinning, Firebase App Check, Geolocation Tracking.', style='List Bullet')
doc.add_paragraph('Visual & Multimedia Design: Photo & Video Editing (Color Grading, Album Design), Print & Digital Media Content creation using Adobe Creative Suite.', style='List Bullet')

doc.add_heading('Projects', level=1)
p1 = doc.add_heading('1. Enterprise Sales Force Automation & POS App (SLFM)', level=2)
p1_sub = doc.add_paragraph('Role: Full-Stack Mobile Developer (Flutter + PHP)')
doc.add_paragraph('Architected an Offline-First Application: Developed a comprehensive field sales app using Flutter to manage attendance, customer billing (POS), stock checking, and damage reporting in low-network areas.', style='List Bullet')
doc.add_paragraph('Developed Custom PHP Backend: Built and secured dynamic PHP endpoints (e.g., attendance.php, index functions) to enforce server-driven rules like dynamic attendance cutoffs and time constraints.', style='List Bullet')
doc.add_paragraph('Engineered a Robust Sync Engine: Implemented a fault-tolerant offline synchronization mechanism using SQLite and Hive to queue local attendance and billing records, resolving complex data duplication issues.', style='List Bullet')
doc.add_paragraph('Implemented Bank-Grade Security: Hardened the app against tampering by integrating FreeRASP for Root/Jailbreak detection, SSL Certificate Pinning, Firebase App Check, and encrypted data handling.', style='List Bullet')
doc.add_paragraph('Repository Management & Security: Sanitized the production codebase for public GitHub release by extracting sensitive PHP/SQL credentials and Firebase configurations into secure, untracked environment variables and utilizing --dart-define for dynamic build injections.', style='List Bullet')
doc.add_paragraph('Background Tracking & Gamification: Configured battery-optimized continuous location tracking. Built a dynamic, animated leaderboard (Top 3 Podium) using Lottie animations and Firebase Realtime Data.', style='List Bullet')

p2 = doc.add_heading('2. Ammu - Mother Care App', level=2)
p2_sub = doc.add_paragraph('Role: Flutter Developer & UI/UX Designer')
doc.add_paragraph('Developed and delivered a user-centric mobile application using Flutter aimed at new mothers, providing health tips, baby care guidance, personalized diet charts, and emergency contacts.', style='List Bullet')
doc.add_paragraph('Led the UI/UX design, focusing on creating intuitive screens and successfully implemented multilingual support (Tamil and English).', style='List Bullet')

doc.add_heading('Professional Experience', level=1)
doc.add_heading('Visual Content Creator | Murugan Digital Studio, Mannargudi', level=2)
doc.add_paragraph('December 2017 – May 2023')
doc.add_paragraph('Managed end-to-end production workflows, including on-location photo/videography, detailed editing, color grading, and final delivery of digital and physical albums/videos for a high volume of events.', style='List Bullet')
doc.add_paragraph('Actively contributed to creating and publishing engaging content for the studio\'s YouTube channel, significantly enhancing its online presence.', style='List Bullet')

doc.add_heading('Education', level=1)
doc.add_paragraph('B.A. (Bachelor of Arts) | Mannai Rajagobalaswamy Government Arts and Science College, Mannargudi, Tamil Nadu (2021 – 2024)')

doc.add_heading('Languages', level=1)
doc.add_paragraph('Tamil: Native / Fluent', style='List Bullet')
doc.add_paragraph('English: Proficient', style='List Bullet')

doc.save(r'C:\Users\LENOVO\Downloads\Updated_CV_Vasanth.docx')
